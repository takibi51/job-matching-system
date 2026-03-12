"""
候補者情報の読み込みモジュール（マルチフォーマット・複数ファイル対応）
- CSV / Excel / テキスト / PDF / 画像(OCR)
- 複数ファイルの統合解析（履歴書 + 職務経歴書 + PF 等）
- 詳細タグ自動付与（職域・スキル・資格・業界・語学 等）
- 個人情報自動除外

個人情報（氏名・連絡先・住所等）は自動除外し、
マッチングに必要な情報のみ抽出する。
"""

import csv
import os
import re
import glob
import io
import unicodedata
from typing import List, Dict, Optional, Tuple

CSV_DIR = os.path.join(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

# ============================================================
# 個人情報フィルタ
# ============================================================

# 個人情報として除外する項目キーワード
PERSONAL_INFO_KEYS = [
    "氏名", "名前", "フルネーム", "本名", "候補者名",
    "電話", "携帯", "TEL", "tel", "Phone", "phone",
    "メール", "メアド", "Email", "email", "e-mail",
    "住所", "自宅", "現住所", "郵便番号", "〒",
    "生年月日", "誕生日", "生まれ",
    "マイナンバー", "保険証", "免許証",
    "LINE", "line", "SNS",
    "家族", "配偶者", "扶養",
]

# テキストから除去する個人情報パターン（正規表現）
_PERSONAL_PATTERNS = [
    r'[\w.+-]+@[\w.-]+\.\w+',                    # メールアドレス
    r'0\d{1,4}[-\s]?\d{1,4}[-\s]?\d{3,4}',      # 電話番号
    r'〒?\d{3}[-\s]?\d{4}',                       # 郵便番号
    r'(?:東京都|北海道|(?:京都|大阪)府|.{2,3}県).{1,6}[市区町村郡].{1,20}[\d丁目番号-]+',  # 住所
]


def _is_personal_info(key: str) -> bool:
    """項目名が個人情報に該当するかチェック"""
    key_lower = key.strip().lower()
    for pi_key in PERSONAL_INFO_KEYS:
        if pi_key.lower() in key_lower:
            return True
    return False


def _remove_personal_from_text(text: str) -> str:
    """テキストから個人情報パターンを除去"""
    cleaned = text
    for pattern in _PERSONAL_PATTERNS:
        cleaned = re.sub(pattern, "[個人情報除外]", cleaned)
    return cleaned


# ============================================================
# ファイル種別判定
# ============================================================

# ファイル名から書類種別を推定
_FILE_TYPE_PATTERNS = {
    "履歴書": ["履歴書", "resume", "rirekisho"],
    "職務経歴書": ["職務経歴", "職歴", "career", "shokumukeireki", "cv"],
    "ポートフォリオ": ["ポートフォリオ", "portfolio", "pf", "作品"],
    "スキルシート": ["スキルシート", "skill", "スキル一覧"],
    "職務要約": ["職務要約", "summary", "要約"],
    "面談メモ": ["面談", "interview", "面接"],
    "推薦状": ["推薦", "recommendation", "紹介"],
    "自己PR": ["自己pr", "自己紹介", "pr"],
}


def _detect_file_type(filename: str, text: str = "") -> str:
    """ファイル名やテキスト内容から書類種別を推定"""
    name_lower = filename.lower()
    for doc_type, patterns in _FILE_TYPE_PATTERNS.items():
        for p in patterns:
            if p.lower() in name_lower:
                return doc_type

    # テキスト内容からも推定
    if text:
        text_head = text[:500].lower()
        if "職務経歴" in text_head or "職歴" in text_head:
            return "職務経歴書"
        if "履歴書" in text_head:
            return "履歴書"
        if "ポートフォリオ" in text_head or "作品" in text_head:
            return "ポートフォリオ"
        if "スキル" in text_head and ("一覧" in text_head or "シート" in text_head):
            return "スキルシート"

    return "その他"


# ============================================================
# マッチング用キーワード抽出（拡張版）
# ============================================================

# 職種・スキル関連のキーワード辞書
_SKILL_KEYWORDS = [
    # デザイン
    "Webデザイナー", "UIデザイナー", "UXデザイナー", "UI/UX",
    "グラフィックデザイナー", "デザイン", "クリエイティブ",
    "Figma", "Photoshop", "Illustrator", "XD", "Sketch", "InDesign",
    "After Effects", "Premiere Pro",
    "アートディレクター", "グラフィックデザイン", "DTP", "印刷",
    "VI", "CI", "UXデザイン", "UIデザイン",
    "GIFアニメーション", "バナー制作",
    # マーケティング
    "Webマーケティング", "デジタルマーケティング", "マーケター",
    "Web広告", "WEB広告", "SNS広告", "Google広告", "Meta広告", "リスティング広告",
    "SEO", "SEM", "LPO", "CRO", "MA", "CRM",
    "コンテンツマーケティング", "SNS運用", "広告運用",
    "ブランディング", "PR", "広報",
    "GEO", "LLMO", "ABテスト", "A/Bテスト", "CVR",
    "GA4", "Google Analytics", "Search Console", "Looker Studio",
    "Ahrefs",
    # ディレクション
    "Webディレクター", "WEBディレクター", "ディレクション", "プロジェクトマネジメント",
    "要件定義", "ワイヤーフレーム", "サイトマップ", "仕様書",
    # 開発・IT
    "エンジニア", "プログラマー", "SE", "フロントエンド", "バックエンド",
    "フルスタック", "インフラ", "SRE", "DevOps",
    "HTML", "CSS", "JavaScript", "TypeScript", "Python", "React", "Vue",
    "Next.js", "Node.js", "Ruby", "Rails", "Go", "Java", "Kotlin", "Swift",
    "PHP", "Laravel", "Django", "Flask",
    "AWS", "GCP", "Azure", "Docker", "Kubernetes",
    "WordPress", "EC", "ECサイト", "Shopify",
    "AI", "機械学習", "データサイエンス", "深層学習",
    # ビジネス
    "営業", "法人営業", "個人営業", "ルート営業", "新規開拓",
    "ソリューション営業", "ソリューションセールス",
    "コンサルタント", "コンサルティング", "コンサルティング営業",
    "プロジェクトマネージャー", "PM", "ディレクター", "プロデューサー",
    "企画", "事業企画", "経営企画", "商品企画", "企画営業",
    "カスタマーサクセス", "CS", "カスタマーサポート",
    "見積作成", "提案書作成", "顧客折衝", "予算管理",
    "チームマネジメント", "メンバー育成", "進捗管理",
    "イベント企画", "展覧会",
    "商品企画", "販売促進", "広告出稿",
    "テレアポ", "DX", "賃貸管理", "仲介",
    # 管理
    "マネジメント", "チームリーダー", "管理職", "事業部長", "部長",
    "人事", "採用", "総務", "経理", "労務", "法務", "財務", "経営管理",
    "情報システム", "社内SE",
    # 医療・ヘルスケア
    "看護師", "准看護師", "保健師", "助産師",
    "医師", "歯科医師", "薬剤師",
    "理学療法士", "作業療法士", "言語聴覚士",
    "臨床検査技師", "臨床工学技士", "放射線技師",
    "管理栄養士", "栄養士", "介護福祉士", "社会福祉士",
    "医療事務", "看護助手", "歯科衛生士", "歯科助手",
    "柔道整復師", "鍼灸師", "あん摩マッサージ指圧師",
    "ケアマネージャー", "介護職",
    # 制作
    "LP制作", "サイト制作", "コーディング", "ライティング",
    "動画制作", "映像制作", "写真撮影", "コピーライター",
    "編集", "校正", "校閲", "原稿作成",
    # ツール
    "Backlog", "Notion", "Slack", "ChatWork", "Microsoft Teams",
    "STORES", "BASE", "Bカート",
    "Salesforce", "HubSpot", "LINE公式アカウント", "LINE公式",
    "Tポイント",
    # EC
    "EC運営", "ECサイト構築", "オンラインショップ", "通販",
    # 分析
    "データ分析", "アクセス解析", "KPI",
    "プロジェクト管理", "業務改善", "BPR",
    "Tableau", "Power BI", "SQL", "Excel VBA",
    # 業界
    "医療", "ヘルスケア", "不動産", "金融", "教育", "EC",
    "BtoB", "BtoC", "SaaS", "IT", "Web",
    "製造", "メーカー", "商社", "小売", "飲食", "物流",
    "建設", "人材", "広告代理店", "コンサル",
]

# 資格辞書
_CERTIFICATIONS = [
    # IT
    "基本情報技術者", "応用情報技術者", "情報処理安全確保支援士",
    "AWS認定", "AWS Certified", "GCP認定",
    "PMP", "ITIL", "CCNA", "CCNP",
    "Oracle", "LPIC", "LinuC",
    "Salesforce", "HubSpot",
    "Google Analytics認定", "Google広告認定",
    "ウェブ解析士", "上級ウェブ解析士",
    # ビジネス
    "MBA", "中小企業診断士", "社会保険労務士", "行政書士",
    "公認会計士", "税理士", "簿記1級", "簿記2級", "簿記3級",
    "FP1級", "FP2級", "FP3級", "ファイナンシャルプランナー",
    "宅地建物取引士", "宅建",
    "TOEIC", "TOEFL", "英検", "IELTS",
    "HSK", "中国語検定",
    "秘書検定", "ビジネス実務法務検定",
    "キャリアコンサルタント",
    # 医療
    "看護師免許", "准看護師免許", "保健師免許", "助産師免許",
    "医師免許", "歯科医師免許", "薬剤師免許",
    "理学療法士免許", "作業療法士免許", "言語聴覚士免許",
    "管理栄養士免許", "介護福祉士", "社会福祉士",
    "ケアマネージャー", "衛生管理者",
    "BLS", "ACLS", "認定看護師", "専門看護師",
    # デザイン
    "色彩検定", "カラーコーディネーター",
    "DTPエキスパート", "Webデザイナー検定",
    # その他
    "普通自動車免許", "大型免許",
    "衛生管理者", "防火管理者",
    "フォークリフト",
    "ネットマーケティング検定",
]

# 業界辞書
_INDUSTRIES = {
    "IT・Web": ["IT", "Web", "SaaS", "テック", "インターネット", "ソフトウェア", "情報通信"],
    "医療・ヘルスケア": ["医療", "病院", "クリニック", "ヘルスケア", "介護", "福祉", "製薬", "医薬"],
    "金融": ["金融", "銀行", "証券", "保険", "ファイナンス", "フィンテック"],
    "不動産": ["不動産", "デベロッパー", "建設", "住宅"],
    "製造・メーカー": ["製造", "メーカー", "工場", "生産", "品質管理"],
    "商社": ["商社", "貿易", "輸出入"],
    "小売・流通": ["小売", "流通", "百貨店", "スーパー", "EC", "通販"],
    "飲食・サービス": ["飲食", "レストラン", "ホテル", "サービス", "接客"],
    "教育": ["教育", "学校", "塾", "EdTech", "研修"],
    "広告・メディア": ["広告", "メディア", "出版", "マスコミ", "PR"],
    "コンサルティング": ["コンサル", "アドバイザリー", "シンクタンク"],
    "人材": ["人材", "派遣", "紹介", "採用", "HR"],
    "物流": ["物流", "運送", "倉庫", "ロジスティクス"],
    "エネルギー": ["エネルギー", "電力", "ガス", "石油"],
    "官公庁・公社": ["官公庁", "公務員", "自治体", "公社"],
}

# 語学辞書
_LANGUAGE_PATTERNS = {
    "英語": [r'英語', r'English', r'TOEIC\s*(\d+)', r'TOEFL\s*(\d+)', r'英検(\d級|準?\d級)',
             r'IELTS\s*([\d.]+)', r'ビジネス英語', r'日常英会話'],
    "中国語": [r'中国語', r'Chinese', r'HSK\s*(\d)', r'中国語検定(\d級)'],
    "韓国語": [r'韓国語', r'Korean', r'TOPIK\s*(\d)', r'ハングル'],
    "フランス語": [r'フランス語', r'French', r'仏語'],
    "スペイン語": [r'スペイン語', r'Spanish', r'西語'],
    "ドイツ語": [r'ドイツ語', r'German', r'独語'],
    "ポルトガル語": [r'ポルトガル語', r'Portuguese'],
}

# 雇用形態
_EMPLOYMENT_TYPES = {
    "正社員": ["正社員", "正規雇用", "常勤"],
    "契約社員": ["契約社員", "有期雇用"],
    "派遣社員": ["派遣", "派遣社員"],
    "パート・アルバイト": ["パート", "アルバイト", "非常勤"],
    "業務委託": ["業務委託", "フリーランス", "個人事業"],
    "役員": ["役員", "取締役", "執行役員"],
}

# 働き方
_WORK_STYLES = {
    "リモート": ["リモート", "在宅", "テレワーク", "remote"],
    "フレックス": ["フレックス", "裁量労働", "flex"],
    "時短": ["時短", "短時間勤務"],
    "副業OK": ["副業", "複業", "ダブルワーク"],
    "転勤なし": ["転勤なし", "転勤不可"],
}

# マネジメント規模パターン
_MGMT_PATTERNS = [
    r'(\d+)\s*名?\s*(?:の\s*)?(?:部下|メンバー|チーム|組織)',
    r'(?:部下|メンバー|チーム)\s*(\d+)\s*名',
    r'(\d+)\s*(?:人|名)\s*規模',
    r'マネジメント\s*[:：]?\s*(\d+)',
]

# 入社可能時期パターン
_AVAILABILITY_PATTERNS = [
    (r'即[日時]', "即日"),
    (r'(\d+)\s*(?:ヶ月|か月|カ月)', None),  # 動的に生成
    (r'(?:来月|翌月)', "1ヶ月後"),
    (r'(?:再来月)', "2ヶ月後"),
    (r'(\d{4})[年/](\d{1,2})月?', None),  # 動的に生成
]

# 転職理由カテゴリ
_CAREER_CHANGE_REASONS = {
    "キャリアアップ": ["キャリアアップ", "スキルアップ", "成長", "ステップアップ", "キャリアチェンジ"],
    "待遇改善": ["年収", "給与", "待遇", "報酬", "昇給"],
    "ワークライフバランス": ["ワークライフ", "残業", "働き方", "ライフワーク", "プライベート"],
    "やりがい": ["やりがい", "挑戦", "新しい", "興味", "志"],
    "人間関係": ["人間関係", "上司", "社風", "文化"],
    "会社都合": ["リストラ", "倒産", "事業縮小", "会社都合", "契約満了"],
    "Uターン・Iターン": ["Uターン", "Iターン", "地元", "地方"],
}

# 部署名 → キーワード マッピング
_DEPT_KEYWORD_MAP = {
    "クリエイティブ": ["Webデザイナー", "デザイン", "クリエイティブ"],
    "UX推進": ["UXデザイン", "UXディレクター", "UI/UX"],
    "マーケティング": ["Webマーケティング", "デジタルマーケティング", "広告運用"],
    "デジタルマーケティング": ["デジタルマーケティング", "Web広告", "マーケティング"],
    "コンサルティング": ["コンサルタント", "コンサルティング営業", "法人営業"],
    "営業": ["営業", "法人営業", "ソリューション営業"],
    "開発": ["エンジニア", "開発", "プログラマー"],
    "企画": ["企画", "事業企画", "プロデューサー"],
    "人事": ["人事", "採用", "HR"],
    "管理": ["管理職", "マネジメント", "部長"],
}

# 役職 → キーワード マッピング
_POSITION_KEYWORD_MAP = {
    "事業部長": ["事業部長", "マネージャー", "管理職", "部長"],
    "チームリーダー": ["チームリーダー", "マネージャー", "リーダー"],
    "マネージャー": ["マネージャー", "管理職", "リーダー"],
    "ディレクター": ["ディレクター", "マネージャー"],
    "リーダー": ["リーダー", "チームリーダー"],
}

# 経験レベル推定用
_EXPERIENCE_LEVEL_MAP = {
    "ジュニア": {"max_years": 3, "keywords": ["未経験", "第二新卒", "新卒", "ジュニア", "アシスタント"]},
    "ミドル": {"max_years": 10, "keywords": ["中堅", "ミドル", "一人称"]},
    "シニア": {"max_years": 20, "keywords": ["シニア", "スペシャリスト", "エキスパート", "主任", "係長"]},
    "リード": {"max_years": 99, "keywords": ["リード", "マネージャー", "課長", "部長", "ディレクター"]},
    "エグゼクティブ": {"max_years": 99, "keywords": ["CTO", "CIO", "VP", "執行役員", "取締役", "役員"]},
}


def _extract_keywords_from_text(text: str) -> List[str]:
    """テキストからマッチング用キーワードを抽出"""
    found = []
    text_lower = text.lower()

    # 短いキーワード（3文字以下の英字）は単語境界でマッチさせる
    _SHORT_WORD_BOUNDARY = re.compile(r'(?<![a-zA-Z]){kw}(?![a-zA-Z])')

    # 長いキーワードから先にマッチさせる（部分一致の誤検出防止）
    sorted_keywords = sorted(_SKILL_KEYWORDS, key=len, reverse=True)
    matched_positions = set()

    for kw in sorted_keywords:
        kw_lower = kw.lower()

        # 短い英字キーワード（SE, PM, CS, MA等）は単語境界で検索
        if len(kw) <= 3 and kw.isascii() and kw.isalpha():
            pattern = re.compile(r'(?<![a-zA-Z])' + re.escape(kw_lower) + r'(?![a-zA-Z])')
            m = pattern.search(text_lower)
            if m:
                pos = m.start()
                kw_range = set(range(pos, pos + len(kw_lower)))
                if not kw_range & matched_positions:
                    found.append(kw)
                    matched_positions.update(kw_range)
        else:
            pos = text_lower.find(kw_lower)
            if pos >= 0:
                kw_range = set(range(pos, pos + len(kw_lower)))
                if not kw_range & matched_positions:
                    found.append(kw)
                    matched_positions.update(kw_range)

    return list(dict.fromkeys(found))  # 重複除去・順序保持


def _extract_certifications(text: str) -> List[str]:
    """テキストから資格を抽出"""
    found = []
    text_lower = text.lower()
    for cert in _CERTIFICATIONS:
        if cert.lower() in text_lower:
            found.append(cert)

    # TOEIC スコア抽出
    toeic_m = re.search(r'TOEIC\s*[:：]?\s*(\d{3,4})\s*点?', text, re.IGNORECASE)
    if toeic_m:
        score = int(toeic_m.group(1))
        found = [f for f in found if "TOEIC" not in f]
        found.append(f"TOEIC {score}点")

    return list(dict.fromkeys(found))


def _extract_industries(text: str) -> List[str]:
    """テキストから経験業界を抽出"""
    found = []
    for industry, keywords in _INDUSTRIES.items():
        for kw in keywords:
            if kw in text:
                found.append(industry)
                break
    return list(dict.fromkeys(found))


def _extract_languages(text: str) -> List[Dict]:
    """テキストから語学力を抽出"""
    results = []
    for lang, patterns in _LANGUAGE_PATTERNS.items():
        for pat in patterns:
            m = re.search(pat, text, re.IGNORECASE)
            if m:
                level = ""
                if m.lastindex and m.group(1):
                    level = m.group(0)
                results.append({"language": lang, "level": level or "あり"})
                break
    return results


def _extract_employment_type(text: str) -> List[str]:
    """テキストから希望雇用形態を抽出"""
    found = []
    for emp_type, keywords in _EMPLOYMENT_TYPES.items():
        for kw in keywords:
            if kw in text:
                found.append(emp_type)
                break
    return list(dict.fromkeys(found))


def _extract_work_styles(text: str) -> List[str]:
    """テキストから希望働き方を抽出"""
    found = []
    for style, keywords in _WORK_STYLES.items():
        for kw in keywords:
            if kw.lower() in text.lower():
                found.append(style)
                break
    return list(dict.fromkeys(found))


def _extract_management_experience(text: str) -> Dict:
    """テキストからマネジメント経験を抽出"""
    result = {"has_experience": False, "team_size": 0, "details": ""}

    # 「マネジメント」「管理」「統括」等のキーワードチェック
    mgmt_keywords = ["マネジメント", "管理", "統括", "部下", "メンバー", "チーム運営", "組織運営"]
    has_mgmt = any(kw in text for kw in mgmt_keywords)

    if has_mgmt:
        result["has_experience"] = True

        # 規模の抽出
        for pat in _MGMT_PATTERNS:
            m = re.search(pat, text)
            if m:
                result["team_size"] = int(m.group(1))
                break

        # 詳細の抽出
        for line in text.split("\n"):
            if any(kw in line for kw in mgmt_keywords):
                result["details"] = line.strip()[:100]
                break

    return result


def _extract_experience_level(text: str, years: int = 0) -> str:
    """経験レベルを推定"""
    # キーワードベースで判定
    for level, config in _EXPERIENCE_LEVEL_MAP.items():
        for kw in config["keywords"]:
            if kw in text:
                return level

    # 経験年数ベースで判定
    if years > 0:
        if years <= 3:
            return "ジュニア"
        elif years <= 7:
            return "ミドル"
        elif years <= 15:
            return "シニア"
        else:
            return "リード"

    return "ミドル"  # デフォルト


def _extract_availability(text: str) -> str:
    """入社可能時期を抽出"""
    for pat, label in _AVAILABILITY_PATTERNS:
        m = re.search(pat, text)
        if m:
            if label:
                return label
            if m.lastindex == 1:
                return f"{m.group(1)}ヶ月後"
            if m.lastindex == 2:
                return f"{m.group(1)}年{m.group(2)}月"
    return ""


def _extract_career_change_reasons(text: str) -> List[str]:
    """転職理由カテゴリを抽出"""
    found = []
    for reason, keywords in _CAREER_CHANGE_REASONS.items():
        for kw in keywords:
            if kw in text:
                found.append(reason)
                break
    return found


def _extract_education(text: str) -> Dict:
    """学歴情報を抽出（個人情報に注意）"""
    result = {"level": "", "field": ""}

    # 最終学歴レベル
    edu_levels = [
        ("博士", ["博士", "PhD", "ドクター"]),
        ("修士", ["修士", "大学院", "MBA", "Master"]),
        ("大卒", ["大学", "学部", "Bachelor"]),
        ("短大・専門", ["短大", "専門学校", "高専"]),
        ("高卒", ["高校", "高等学校"]),
    ]
    for level, keywords in edu_levels:
        if any(kw in text for kw in keywords):
            result["level"] = level
            break

    # 専攻分野
    fields = [
        ("情報工学", ["情報", "コンピュータ", "IT"]),
        ("経営学", ["経営", "商学", "ビジネス", "MBA"]),
        ("経済学", ["経済"]),
        ("法学", ["法学", "法律"]),
        ("医学・看護学", ["医学", "看護", "薬学", "医療"]),
        ("工学", ["工学", "理工", "機械"]),
        ("文学", ["文学", "人文", "社会学"]),
        ("教育学", ["教育"]),
        ("デザイン・芸術", ["デザイン", "芸術", "美術"]),
    ]
    for field, keywords in fields:
        if any(kw in text for kw in keywords):
            result["field"] = field
            break

    return result


def _extract_age(text: str) -> int:
    """テキストから年齢を抽出（複数パターン対応）"""
    # 「年齢: 30歳」パターン
    m = re.search(r'(?:年齢|Age)[：:\s]*(\d{1,2})歳?', text)
    if m:
        return int(m.group(1))
    # 「満50歳」「満 50 歳」パターン
    m = re.search(r'満\s*(\d+)\s*歳', text)
    if m:
        age = int(m.group(1))
        if 18 <= age <= 70:
            return age
    # 「1976年01月23日生」→ 生年月日から年齢を計算（基準年: 2026）
    m = re.search(r'(\d{4})年\s*(\d{1,2})月\s*(\d{1,2})日生', text)
    if m:
        birth_year = int(m.group(1))
        birth_month = int(m.group(2))
        birth_day = int(m.group(3))
        # 2026年3月12日を基準に計算
        age = 2026 - birth_year
        # 誕生日がまだ来ていなければ1歳引く
        if (birth_month, birth_day) > (3, 12):
            age -= 1
        if 18 <= age <= 70:
            return age
    # 「30歳」パターン（一般的な記載）
    m = re.search(r'(\d{2})歳', text)
    if m:
        age = int(m.group(1))
        if 18 <= age <= 70:
            return age
    return 0


def _extract_gender(text: str) -> str:
    """テキストから性別を抽出"""
    # 「性別：女」「性別: 男」パターン
    m = re.search(r'性別[：:\s]*([男女])', text)
    if m:
        return m.group(1) + "性"
    # テキスト中の単独の「男」「女」（性別文脈で出現するもの）
    # 「日生（満50歳）」の近くや履歴書の基本情報欄にある性別表記
    m = re.search(r'歳\s*[）\)]\s*([男女])\b', text)
    if m:
        return m.group(1) + "性"
    # 行頭や区切り文字の後に単独で出現する「男」「女」
    m = re.search(r'(?:^|[\s　\t])\s*([男女])\s*(?:$|[\s　\t\n])', text, re.MULTILINE)
    if m:
        return m.group(1) + "性"
    return ""


def _extract_career_summary(text: str) -> str:
    """
    職務経歴要約・経歴略歴を抽出（最大800文字）。
    職務経歴要約セクションがあればそれを優先、なければ履歴書の職歴タイムラインを構築。
    さらに、CV中の会社名+期間パターンからも簡易タイムラインを生成。
    """
    # 職務経歴要約セクションを探す
    summary_headers = [
        r'■\s*職務経歴要約',
        r'■\s*職務要約',
        r'■\s*経歴要約',
        r'Career\s*Summary',
        r'職務経歴要約',
        r'職務要約',
        r'経歴要約',
    ]
    best_content = ""
    for header_pat in summary_headers:
        m = re.search(header_pat + r'[：:\s]*\n?([\s\S]*?)(?=\n\s*■|\n\s*\n\s*\n)', text)
        if m:
            content = m.group(1).strip()
            if content and len(content) > len(best_content):
                best_content = content
        if not best_content:
            m = re.search(header_pat + r'[：:\s]*\n?([\s\S]*?)$', text)
            if m:
                content = m.group(1).strip()
                if content and len(content) > len(best_content):
                    best_content = content
    if best_content:
        return best_content[:800]

    # 履歴書の職歴テーブルから簡易タイムラインを構築
    career_entries = _extract_resume_career_history(text)
    if career_entries:
        timeline = " → ".join(career_entries[:5])
        return timeline[:800]

    # CV中の「会社名 年月〜年月」パターンから抽出
    _SKIP_WORDS = {"期間", "業務内容", "項目", "内容", "年月", "備考", "使用ツール", "概要",
                   "営業", "管理", "企画", "開発", "エリア", "担当", "マネージャー"}
    cv_company_pattern = re.compile(
        r'((?:株式会社|有限会社|合同会社|医療法人)?[^\n]{2,25}'
        r'(?:株式会社|有限会社|合同会社)?)'
        r'\s+(\d{4})年\s*(\d{1,2})月\s*[〜~～\-–]'
    )
    cv_entries = []
    for m in cv_company_pattern.finditer(text):
        company = m.group(1).strip()
        company = re.sub(r'（[^）]*）', '', company).strip()
        # テーブルヘッダー等のスキップ（部分一致）
        if not company or len(company) < 2:
            continue
        if any(sw in company for sw in _SKIP_WORDS):
            continue
        cv_entries.append(f"{company} ({m.group(2)}年{m.group(3)}月〜)")
    if cv_entries:
        timeline = " → ".join(dict.fromkeys(cv_entries).keys())
        return timeline[:800]

    return ""


def _extract_self_pr(text: str) -> str:
    """
    自己PRセクションを抽出（最大800文字）。
    セクションヘッダーから次の■セクションまで、または末尾まで取得。
    テンプレートの列ヘッダー（「志望動機、アピールポイント、特技など」）はスキップ。
    """
    # テンプレート列ヘッダーのパターン（履歴書の欄タイトル）
    _TEMPLATE_HEADERS = re.compile(
        r'志望動機[、,・]\s*(?:自己PR[、,・]\s*)?アピールポイント'
        r'|アピールポイント[、,・]\s*特技'
        r'|アピールポイントなど'
        r'|特技など'
    )

    pr_headers = [
        r'■\s*自己\s*PR',
        r'自己\s*PR',
        r'■\s*志望の動機',
        r'■\s*アピールポイント',
    ]
    best_content = ""
    for header_pat in pr_headers:
        for m in re.finditer(header_pat + r'[：:\s]*\n?([\s\S]*?)(?=\n\s*■|\n\s*\n\s*\n)', text):
            # テンプレート欄タイトルの一部としてマッチした場合はスキップ
            match_start = max(0, m.start() - 20)
            context = text[match_start:m.start() + 40]
            if _TEMPLATE_HEADERS.search(context):
                continue
            content = m.group(1).strip()
            if content and len(content) > 20 and len(content) > len(best_content):
                best_content = content
        # ■がない場合（末尾まで）
        if not best_content:
            m = re.search(header_pat + r'[：:\s]*\n?([\s\S]*?)$', text)
            if m:
                match_start = max(0, m.start() - 20)
                context = text[match_start:m.start() + 40]
                if _TEMPLATE_HEADERS.search(context):
                    continue
                content = m.group(1).strip()
                if content and len(content) > 20 and len(content) > len(best_content):
                    best_content = content
    if best_content:
        return best_content[:800]

    # フォールバック1: CV中の人物特性・強みを表す記述を汎用的に探す
    _person_trait_keywords = [
        # コミュニケーション系
        "プレゼン", "伝わる", "伝える", "説明力", "折衝", "交渉",
        "コミュニケーション", "ヒアリング", "傾聴", "対話",
        # 行動特性系
        "素早く対応", "行動を起こ", "行動力", "フットワーク", "レスポンス",
        "スピード", "迅速", "率先",
        # 対人スキル系
        "真摯に", "誠実", "信頼", "丁寧", "親身",
        "顧客に提供", "最善の対応", "課題解決",
        # 知識・スキル向上系
        "勉強を行い", "自己研鑽", "スキルアップ", "資格を取得", "獲得し",
        "学び", "習得",
        # リーダーシップ系
        "部下", "メンバー", "チームを", "組織を", "牽引", "統率",
        "活気づけ", "マネジメント",
        # 成果・実績系
        "達成", "貢献", "実績", "受賞", "MVP", "改善",
        "売上", "案件を獲得", "目標", "V字回復",
        # 思考力系
        "戦略的", "論理的", "分析", "企画力", "発想力", "創意工夫",
        "提案力", "問題解決",
        # 姿勢系
        "粘り強", "コミット", "挑戦", "主体的", "積極的", "柔軟",
        "最後まで", "やり遂げ",
    ]
    lines = text.split("\n")
    pr_lines = []
    seen = set()
    for line in lines:
        line = line.strip()
        # 短すぎる行や見出し行はスキップ
        if not line or len(line) < 15 or line in seen:
            continue
        if any(kw in line for kw in _person_trait_keywords):
            pr_lines.append(line)
            seen.add(line)
    if pr_lines:
        return "\n".join(pr_lines[:10])[:800]

    # フォールバック2: 「強み」セクション内の記述文を探す
    strength_section = re.search(
        r'(?:強み|得意|特徴|長所)[：:\s]*\n?([\s\S]*?)(?=\n\s*■|\n\s*\n\s*\n|$)',
        text
    )
    if strength_section:
        content = strength_section.group(1).strip()
        if content and len(content) > 20:
            return content[:800]

    return ""


def _extract_career_highlights(text: str) -> str:
    """職務経歴の成果・活かせる経験・業務実績から人物面の補足情報を抽出"""
    highlights = []
    seen = set()

    # 1. 「活かせる経験、知識」「活かせるスキル」セクション
    exp_headers = [
        r'■\s*活かせる経験[、,]?\s*(?:知識|スキル)?',
        r'■\s*活かせるスキル',
        r'■\s*スキル・経験',
    ]
    for header_pat in exp_headers:
        m = re.search(header_pat + r'[：:\s]*\n?([\s\S]*?)(?=\n\s*■)', text)
        if m:
            content = m.group(1).strip()
            if content and len(content) > 20 and content not in seen:
                highlights.append(content[:400])
                seen.add(content)
            break

    # 2. 【成果】セクションから実績を収集
    achievement_patterns = [
        r'【成果】\s*\n?((?:[-・]\s*.+\n?)+)',
        r'【実績】\s*\n?((?:[-・]\s*.+\n?)+)',
    ]
    for pat in achievement_patterns:
        for m in re.finditer(pat, text):
            content = m.group(1).strip()
            if content and content not in seen:
                highlights.append(content[:200])
                seen.add(content)

    # 3. 【活かしたスキル】から人物特性の記述を収集
    skill_sections = re.finditer(
        r'【活かした(?:スキル|経験[・]?スキル)】\s*\n?((?:[-・]\s*.+\n?)+)', text
    )
    for m in skill_sections:
        content = m.group(1).strip()
        # 長い説明文がある行のみ（人物面の記述）
        for line in content.split("\n"):
            line = re.sub(r'^[-・]\s*', '', line).strip()
            if len(line) > 30 and line not in seen:
                highlights.append(line[:200])
                seen.add(line)

    if not highlights:
        return ""
    return "\n".join(highlights)[:600]


def _extract_tools_from_text(text: str) -> List[str]:
    """使用ツールセクションからツール名を抽出"""
    tools = []
    # 「使用ツール」「ツール」「使用技術」セクションを探す
    tool_headers = [
        r'使用ツール[：:\s]*',
        r'ツール[：:\s]*',
        r'使用技術[：:\s]*',
        r'使用ソフト[：:\s]*',
        r'開発環境[：:\s]*',
    ]
    for header_pat in tool_headers:
        m = re.search(header_pat + r'(.+?)(?:\n\n|\n(?=[■●▪])|$)', text, re.DOTALL)
        if m:
            tool_text = m.group(1).strip()
            # カンマ、スラッシュ、改行で分割
            parts = re.split(r'[,、/／\n・]+', tool_text)
            for part in parts:
                part = part.strip()
                if part and 1 < len(part) < 40:
                    tools.append(part)

    # 既知のツール名でもテキスト全体からマッチ
    known_tools = [
        "Figma", "Illustrator", "Photoshop", "XD", "Sketch", "InDesign",
        "After Effects", "Premiere Pro", "Canva",
        "Backlog", "Notion", "Slack", "ChatWork", "Microsoft Teams",
        "STORES", "BASE", "Bカート", "Shopify",
        "WordPress", "Webflow",
        "GA4", "Google Analytics", "Search Console", "Looker Studio",
        "Ahrefs", "Tableau", "Power BI",
        "Salesforce", "HubSpot",
    ]
    text_lower = text.lower()
    for tool in known_tools:
        if tool.lower() in text_lower and tool not in tools:
            tools.append(tool)

    return list(dict.fromkeys(tools))


def _wareki_to_seireki(era: str, year_num: int) -> int:
    """和暦を西暦に変換"""
    era_map = {
        "令和": 2018, "平成": 1988, "昭和": 1925, "大正": 1911, "明治": 1867,
    }
    base = era_map.get(era, 0)
    return base + year_num if base else 0


def _extract_resume_career_history(text: str) -> List[str]:
    """
    履歴書の学歴・職歴テーブルから職歴エントリを抽出。
    西暦（2025 5）と和暦（平成19 4）の両方に対応。
    「入社」「設立」等のアクションを含むエントリのみ抽出（学歴・資格は除外）。
    """
    entries = []

    # 和暦パターン: 平成19 4 A株式会社 入社
    wareki_pattern = re.compile(
        r'(平成|令和|昭和)\s*(\d{1,2})\s+(\d{1,2})\s+'
        r'((?:株式会社|有限会社|合同会社|医療法人|社会福祉法人)?[^\n入設退転]{2,30}'
        r'(?:株式会社|有限会社|合同会社)?)'
        r'(?:（[^）]*）)?\s*'
        r'(?:に?\s*)?(入社|設立|を設立|退社|退職|転籍)'
    )
    # 西暦パターン: 2025 5 A株式会社 入社
    seireki_pattern = re.compile(
        r'(\d{4})\s+(\d{1,2})\s+'
        r'((?:株式会社|有限会社|合同会社|医療法人|社会福祉法人)?[^\n入設退転]{2,30}'
        r'(?:株式会社|有限会社|合同会社)?)'
        r'(?:（[^）]*）)?\s*'
        r'(?:に?\s*)?(入社|設立|を設立|退社|退職|転籍)'
    )
    is_current_pattern = re.compile(r'現在に至る|在職中|現職|退職予定')

    # 統一フォーマット: (西暦year, month, company, action, match_end)
    raw_entries = []

    for m in wareki_pattern.finditer(text):
        era = m.group(1)
        era_year = int(m.group(2))
        year = _wareki_to_seireki(era, era_year)
        month = m.group(3)
        company = m.group(4).strip()
        action = m.group(5) or ""
        raw_entries.append((year, month, company, action, m.end()))

    if not raw_entries:
        for m in seireki_pattern.finditer(text):
            year = int(m.group(1))
            month = m.group(2)
            company = m.group(3).strip()
            action = m.group(4) or ""
            raw_entries.append((year, month, company, action, m.end()))

    for i, (year, month, company, action, end_pos) in enumerate(raw_entries):
        if action in ("退社", "退職"):
            continue

        company_clean = re.sub(r'（[^）]*）', '', company).strip()
        if not company_clean or len(company_clean) < 2:
            continue

        end_str = "現在"
        remaining = text[end_pos:]
        if i < len(raw_entries) - 1:
            next_year, next_month, _, next_action, _ = raw_entries[i + 1]
            if next_action in ("退社", "退職"):
                end_str = f"{next_year}年{next_month}月"
            elif not is_current_pattern.search(remaining[:200]):
                end_str = f"{next_year}年{next_month}月頃"
        else:
            if is_current_pattern.search(remaining[:300]):
                end_str = "現在"

        entry = f"{company_clean} ({year}年{month}月〜{end_str})"
        entries.append(entry)

    # 重複除去（同じ会社名の重複エントリを排除）
    return list(dict.fromkeys(entries))


def _extract_experience_years(text: str) -> int:
    """テキストから経験年数を抽出"""
    # 「経験年数: 5年」パターン
    m = re.search(r'(?:経験年数|経験)[：:\s]*(\d{1,2})\s*年', text)
    if m:
        return int(m.group(1))
    # 「5年以上の経験」パターン
    m = re.search(r'(\d{1,2})\s*年\s*(?:以上)?(?:の)?(?:経験|実務)', text)
    if m:
        return int(m.group(1))
    return 0


def _extract_salary(text: str) -> Tuple[int, int]:
    """テキストから年収を抽出して（min, max）を返す"""
    # 「現年収: 380万円」「年収380万」パターン
    m = re.search(r'(?:現年収|年収|想定年収|希望年収)[：:\s]*(\d{2,4})万', text)
    if m:
        val = int(m.group(1))
        return (int(val * 0.9), int(val * 1.5))
    # 「300万〜500万」パターン
    m = re.search(r'(\d{2,4})万[円〜~～-]+(\d{2,4})万', text)
    if m:
        return (int(m.group(1)), int(m.group(2)))
    return (0, 0)


def _extract_location(text: str) -> str:
    """テキストから希望勤務地を抽出"""
    m = re.search(r'(?:希望勤務地|勤務地|勤務希望)[：:\s]*(.+?)(?:\n|$|、|。)', text)
    if m:
        return m.group(1).strip()[:10]
    # 関西圏のキーワードがあれば
    kansai = ["大阪", "京都", "神戸", "兵庫", "奈良"]
    for loc in kansai:
        if loc in text:
            return loc
    return "大阪"


def _extract_job_history(text: str) -> List[Dict]:
    """テキストから職歴を抽出"""
    histories = []
    # 「20XX年〜」「20XX年X月〜20XX年X月」パターン
    blocks = re.split(r'\n(?=\d{4}[年/])', text)
    for block in blocks:
        m = re.match(r'(\d{4})[年/]\s*(\d{1,2})?月?\s*[〜~～\-–]\s*(?:(\d{4})[年/]?\s*(\d{1,2})?月?|現在|在職中)?', block)
        if m:
            entry = {
                "period": m.group(0).strip(),
                "content": block[m.end():].strip()[:200]
            }
            # 会社名抽出
            company_m = re.search(r'(?:株式会社|有限会社|合同会社)?[^\n]{2,20}(?:株式会社|有限会社|合同会社|病院|クリニック)?', block[m.end():])
            if company_m:
                entry["company"] = company_m.group(0).strip()[:30]
            histories.append(entry)

    return histories[:10]  # 最大10社


def _extract_interview_tags(text: str) -> List[str]:
    """面談・人物に関するタグを抽出"""
    tags = []
    tag_patterns = {
        "論理的": ["論理", "ロジカル", "分析的"],
        "コミュニケーション力": ["コミュニケーション", "対人", "折衝", "交渉力"],
        "リーダーシップ": ["リーダー", "統率", "牽引"],
        "主体的": ["主体", "自発的", "積極的", "自走"],
        "協調性": ["協調", "チームワーク", "協力的"],
        "粘り強い": ["粘り強", "コミット", "やり遂げ"],
        "成長志向": ["成長", "学習意欲", "向上心", "挑戦"],
        "柔軟性": ["柔軟", "適応力", "臨機応変"],
        "誠実": ["誠実", "真面目", "責任感"],
        "行動力": ["行動力", "フットワーク", "スピード感"],
        "創造的": ["創造", "クリエイティブ", "発想力", "企画力"],
        "ストレス耐性": ["ストレス耐性", "タフ", "忍耐"],
        "細部への注意": ["細かい", "丁寧", "正確", "几帳面"],
        "プレゼン力": ["プレゼン", "説明力", "伝える力"],
        "問題解決力": ["問題解決", "課題解決", "改善"],
    }
    for tag, keywords in tag_patterns.items():
        for kw in keywords:
            if kw in text:
                tags.append(tag)
                break
    return tags


def _extract_project_achievements(text: str) -> List[str]:
    """実績・成果を抽出"""
    achievements = []
    # 数値を含む成果文
    patterns = [
        r'(?:売上|利益|コスト).{0,20}(?:\d+[%％万億])',
        r'(?:前年比|対前年|YoY).{0,10}\d+[%％]',
        r'(?:\d+[%％]).{0,10}(?:向上|改善|達成|削減|増加)',
        r'(?:\d+万?円?).{0,10}(?:売上|受注|契約)',
        r'(?:新規|獲得).{0,20}\d+(?:件|社|名)',
    ]
    for pat in patterns:
        matches = re.findall(pat, text)
        for m in matches:
            if m.strip() and len(m.strip()) > 5:
                achievements.append(m.strip()[:80])

    # 箇条書きの成果
    for line in text.split("\n"):
        line = line.strip()
        if re.match(r'^[-・●▪■]\s*', line):
            cleaned = re.sub(r'^[-・●▪■]\s*', '', line)
            if re.search(r'\d', cleaned) and len(cleaned) > 10:
                achievements.append(cleaned[:80])

    return list(dict.fromkeys(achievements))[:10]


# ============================================================
# 拡張タグ抽出（全カテゴリ統合）
# ============================================================

def extract_all_tags(text: str, info: Dict = None) -> Dict:
    """テキストから全カテゴリのタグを抽出"""
    if info is None:
        info = {}

    # 経験年数
    exp_years = _extract_experience_years(text)
    exp_str = info.get("経験年数", "")
    if not exp_years and exp_str:
        m = re.search(r'(\d+)', exp_str)
        if m:
            exp_years = int(m.group(1))

    tags = {
        "skills": _extract_keywords_from_text(text),
        "certifications": _extract_certifications(text),
        "industries": _extract_industries(text),
        "languages": _extract_languages(text),
        "employment_type": _extract_employment_type(text),
        "work_styles": _extract_work_styles(text),
        "management": _extract_management_experience(text),
        "experience_level": _extract_experience_level(text, exp_years),
        "experience_years": exp_years,
        "availability": _extract_availability(text),
        "career_change_reasons": _extract_career_change_reasons(text),
        "education": _extract_education(text),
        "interview_tags": _extract_interview_tags(text),
        "achievements": _extract_project_achievements(text),
        "job_history_count": len(_extract_job_history(text)),
    }

    return tags


# ============================================================
# 検索条件の構築（拡張版）
# ============================================================

def _build_conditions(info: Dict, strengths: List[Tuple[str, str]],
                      full_text: str, tags: Dict = None) -> Dict:
    """候補者データから検索条件を自動構築"""
    conditions = {
        "keywords": [],
        "location": "大阪",
        "salary_min": 0,
        "salary_max": 0,
        "age": 0,
        "prefer_kansai": True,
        "extra_keywords": [],
    }

    # --- info辞書からの抽出 ---
    # 年齢
    age_str = info.get("年齢", "")
    age_m = re.search(r'(\d+)', age_str)
    if age_m:
        conditions["age"] = int(age_m.group(1))

    # 年収（現年収を保存 + 希望年収レンジを算出）
    salary_str = info.get("現年収", "") or info.get("年収", "") or info.get("希望年収", "")
    salary_nums = re.findall(r'[\d,]+', salary_str.replace(",", ""))
    if salary_nums:
        val = int(salary_nums[0])
        conditions["current_salary"] = val  # 現年収を保存
        conditions["salary_min"] = int(val * 0.9)
        conditions["salary_max"] = int(val * 1.5)

    # 役割・部署・役職
    role = info.get("役割", "") or info.get("職種", "") or info.get("希望職種", "")
    department = info.get("所属部署", "") or info.get("部署", "")
    position = info.get("役職", "") or info.get("ポジション", "")

    if role:
        conditions["keywords"].append(role)

    for dept_key, kws in _DEPT_KEYWORD_MAP.items():
        if dept_key in department:
            conditions["keywords"].extend(kws)
            break

    if position and position != "一般":
        for pos_key, kws in _POSITION_KEYWORD_MAP.items():
            if pos_key in position:
                conditions["keywords"].extend(kws)
                break

    # 勤務地
    loc = info.get("希望勤務地", "") or info.get("勤務地", "")
    if loc:
        conditions["location"] = loc

    # --- タグからの補完 ---
    if tags:
        # スキルタグからキーワード追加
        for skill in tags.get("skills", []):
            if skill not in conditions["keywords"]:
                conditions["keywords"].append(skill)
        # 業界からキーワード追加
        for ind in tags.get("industries", []):
            if ind not in conditions["extra_keywords"]:
                conditions["extra_keywords"].append(ind)
        # 資格からキーワード追加
        for cert in tags.get("certifications", [])[:3]:
            if cert not in conditions["extra_keywords"]:
                conditions["extra_keywords"].append(cert)

    # --- 強みからの抽出 ---
    strength_text = " ".join(f"{n} {d}" for n, d in strengths)
    str_keywords = _extract_keywords_from_text(strength_text)
    for kw in str_keywords:
        if kw not in conditions["keywords"] and kw not in conditions["extra_keywords"]:
            conditions["extra_keywords"].append(kw)

    # --- full_text からの補完 ---
    if not conditions["age"]:
        conditions["age"] = _extract_age(full_text)
    if not conditions.get("current_salary") and not conditions["salary_min"]:
        conditions["salary_min"], conditions["salary_max"] = _extract_salary(full_text)
    elif not conditions["salary_min"]:
        conditions["salary_min"], conditions["salary_max"] = _extract_salary(full_text)
    if not conditions["keywords"]:
        conditions["keywords"] = _extract_keywords_from_text(full_text)[:5]
    if not conditions["extra_keywords"]:
        conditions["extra_keywords"] = _extract_keywords_from_text(full_text)[5:10]

    # 勤務地の補完
    if conditions["location"] == "大阪":
        extracted_loc = _extract_location(full_text)
        if extracted_loc != "大阪":
            conditions["location"] = extracted_loc

    # 重複除去・制限
    conditions["keywords"] = list(dict.fromkeys(conditions["keywords"]))[:10]
    conditions["extra_keywords"] = [
        kw for kw in dict.fromkeys(conditions["extra_keywords"])
        if kw not in conditions["keywords"]
    ][:10]

    return conditions


# ============================================================
# CSV 読み込み（既存フォーマット互換）
# ============================================================

def load_candidate_csv(filepath: str) -> Optional[Dict]:
    """候補者CSVを読み込み、検索条件に変換"""
    if not os.path.exists(filepath):
        return None

    rows = []
    with open(filepath, "r", encoding="utf-8") as f:
        reader = csv.reader(f)
        for row in reader:
            rows.append(row)

    candidate = {
        "name": os.path.basename(filepath),
        "info": {},
        "strengths": [],
        "keywords": [],
        "extra_keywords": [],
    }

    section = None
    full_lines = []

    for row in rows:
        if not row or all(cell.strip() == "" for cell in row):
            continue

        first_cell = row[0].strip()
        full_lines.append(" ".join(cell.strip() for cell in row if cell.strip()))

        if first_cell == "候補者情報" or first_cell == "項目":
            section = "info"
            continue
        elif first_cell in ("候補者の強み", "強み"):
            section = "strengths"
            continue
        elif first_cell in ("転職先候補求人リスト", "No."):
            section = "jobs"
            continue

        if section == "info" and len(row) >= 2:
            key = row[0].strip()
            val = row[1].strip()
            if key and val and not _is_personal_info(key):
                candidate["info"][key] = val
        elif section == "strengths" and len(row) >= 2:
            strength_name = row[0].strip()
            strength_detail = row[1].strip() if len(row) > 1 else ""
            if strength_name:
                candidate["strengths"].append((strength_name, strength_detail))

    full_text = "\n".join(full_lines)

    # --- CSV候補者でも人物要約・経歴略歴を生成 ---
    info = candidate["info"]

    # 経歴略歴: 入社日・現職企業・所属部署・役割から構築
    if "経歴略歴" not in info:
        career_parts = []
        if info.get("現職企業"):
            career_parts.append(info["現職企業"])
        if info.get("所属部署"):
            career_parts.append(f"{info['所属部署']}所属")
        if info.get("役割") or info.get("役職"):
            role = info.get("役割", "") or info.get("役職", "")
            if role and role != "一般":
                career_parts.append(role)
        if info.get("入社日"):
            career_parts.append(f"（{info['入社日']}）")
        if career_parts:
            info["経歴略歴"] = " ".join(career_parts)

    # 人物要約: 強みの詳細テキストから構築
    if "人物要約" not in info and candidate["strengths"]:
        summary_parts = []
        for name, detail in candidate["strengths"][:5]:
            if detail and len(detail) > 10:
                summary_parts.append(f"・{name}: {detail[:80]}")
            elif name:
                summary_parts.append(f"・{name}")
        if summary_parts:
            info["人物要約"] = "\n".join(summary_parts)

    # full_text からも年齢・性別を補完
    if "年齢" not in info:
        age_val = _extract_age(full_text)
        if age_val:
            info["年齢"] = f"{age_val}歳"
    if "性別" not in info:
        gender_val = _extract_gender(full_text)
        if gender_val:
            info["性別"] = gender_val

    tags = extract_all_tags(full_text, info)
    candidate["tags"] = tags
    candidate["conditions"] = _build_conditions(
        info, candidate["strengths"], full_text, tags
    )
    return candidate


# ============================================================
# テキストファイル読み込み
# ============================================================

def load_candidate_text(text: str, filename: str = "テキスト入力") -> Optional[Dict]:
    """プレーンテキストから候補者情報を抽出"""
    if not text or not text.strip():
        return None

    cleaned = _remove_personal_from_text(text)

    candidate = {
        "name": filename,
        "info": {},
        "strengths": [],
        "keywords": [],
        "extra_keywords": [],
    }

    # 「項目: 値」パターンで info を抽出
    for line in cleaned.split("\n"):
        line = line.strip()
        if not line:
            continue
        # 箇条書き行はスキップ
        if line.startswith(("・", "●", "▪", "■", "【", "-")):
            continue
        # 「項目：値」「項目: 値」パターン（キー名は2〜15文字）
        m = re.match(r'^([^：:]{2,15})[：:](.+)$', line)
        if m:
            key = m.group(1).strip()
            val = m.group(2).strip()
            if key and val and not _is_personal_info(key) and len(candidate["info"]) < 30:
                candidate["info"][key] = val

    # 「強み」「スキル」「経験」セクションの抽出
    strength_section = re.search(
        r'(?:強み|スキル|得意|経験|実績)[：:\s]*\n?((?:[-・●▪■]\s*.+\n?)+)',
        cleaned, re.MULTILINE
    )
    if strength_section:
        for line in strength_section.group(1).split("\n"):
            line = re.sub(r'^[-・●▪■]\s*', '', line).strip()
            if line:
                candidate["strengths"].append((line, ""))

    # 職歴セクション抽出
    career_section = re.search(
        r'(?:職歴|職務経歴|経歴|Career)[：:\s]*\n((?:.+\n?)+)',
        cleaned, re.MULTILINE | re.IGNORECASE
    )
    if career_section:
        career_text = career_section.group(1)
        histories = _extract_job_history(career_text)
        if histories:
            candidate["info"]["職歴概要"] = f"{len(histories)}社経験"

    # --- 拡張抽出: 年齢・性別・経歴略歴・人物要約 ---
    # 年齢（文字列として保存、既存値を上書きしない）
    if "年齢" not in candidate["info"]:
        age_val = _extract_age(cleaned)
        if age_val:
            candidate["info"]["年齢"] = f"{age_val}歳"

    # 性別
    if "性別" not in candidate["info"]:
        gender_val = _extract_gender(cleaned)
        if gender_val:
            candidate["info"]["性別"] = gender_val

    # 経歴略歴（職務経歴要約 or 履歴書の職歴タイムライン）
    if "経歴略歴" not in candidate["info"]:
        career_summary_val = _extract_career_summary(cleaned)
        if career_summary_val:
            candidate["info"]["経歴略歴"] = career_summary_val

    # 人物要約（自己PR + 職務経歴の成果・業務内容から補完）
    if "人物要約" not in candidate["info"]:
        self_pr_val = _extract_self_pr(cleaned)
        if self_pr_val:
            candidate["info"]["人物要約"] = self_pr_val

    # 人物要約が300字未満の場合、職務経歴の成果・業務実績で補完
    current_pr = candidate["info"].get("人物要約", "")
    if len(current_pr) < 300:
        supplements = _extract_career_highlights(cleaned)
        if supplements:
            if current_pr:
                combined = current_pr + "\n" + supplements
            else:
                combined = supplements
            candidate["info"]["人物要約"] = combined[:800]

    # 使用ツール抽出（タグに追加）
    tools_from_text = _extract_tools_from_text(cleaned)

    tags = extract_all_tags(cleaned, candidate["info"])

    # 使用ツールをスキルタグに統合（重複除去）
    if tools_from_text:
        existing_skills = set(tags.get("skills", []))
        for tool in tools_from_text:
            if tool not in existing_skills:
                tags.setdefault("skills", []).append(tool)
                existing_skills.add(tool)

    candidate["tags"] = tags

    # candidateトップレベルの keywords/extra_keywords にもタグから反映
    candidate["keywords"] = list(tags.get("skills", []))
    candidate["extra_keywords"] = list(tags.get("industries", []))

    candidate["conditions"] = _build_conditions(
        candidate["info"], candidate["strengths"], cleaned, tags
    )
    return candidate


# ============================================================
# PDF 読み込み
# ============================================================

def load_candidate_pdf(filepath_or_bytes, filename: str = "PDF") -> Optional[Dict]:
    """PDFから候補者情報を抽出"""
    try:
        import pdfplumber
    except ImportError:
        return None

    import logging
    logging.getLogger("pdfminer").setLevel(logging.ERROR)

    text_parts = []
    try:
        if isinstance(filepath_or_bytes, (str, os.PathLike)):
            pdf = pdfplumber.open(filepath_or_bytes)
        else:
            pdf = pdfplumber.open(io.BytesIO(filepath_or_bytes))

        for page in pdf.pages[:20]:  # 最大20ページ
            try:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            except Exception:
                pass

            # テーブルも抽出
            try:
                tables = page.extract_tables()
                for table in tables:
                    for row in table:
                        if row:
                            cells = [str(c).strip() for c in row if c]
                            if cells:
                                text_parts.append(" ".join(cells))
            except Exception:
                pass
        pdf.close()
    except Exception:
        return None

    full_text = "\n".join(text_parts)
    if not full_text.strip():
        return None

    # CJK互換文字等を正規化（⾃⼰→自己 等）
    full_text = unicodedata.normalize("NFKC", full_text)

    return load_candidate_text(full_text, filename)


# ============================================================
# 画像 読み込み（OCR）
# ============================================================

def load_candidate_image(filepath_or_bytes, filename: str = "画像") -> Optional[Dict]:
    """画像からOCRで候補者情報を抽出"""
    try:
        import pytesseract
        from PIL import Image
    except ImportError:
        return None

    try:
        if isinstance(filepath_or_bytes, (str, os.PathLike)):
            img = Image.open(filepath_or_bytes)
        else:
            img = Image.open(io.BytesIO(filepath_or_bytes))

        # OCR実行（日本語 + 英語）
        text = pytesseract.image_to_string(img, lang="jpn+eng")
    except Exception:
        return None

    if not text or not text.strip():
        return None

    text = unicodedata.normalize("NFKC", text)
    return load_candidate_text(text, filename)


# ============================================================
# Excel 読み込み
# ============================================================

def load_candidate_excel(filepath_or_bytes, filename: str = "Excel") -> Optional[Dict]:
    """Excelから候補者情報を抽出"""
    try:
        import pandas as pd
    except ImportError:
        return None

    try:
        if isinstance(filepath_or_bytes, (str, os.PathLike)):
            df = pd.read_excel(filepath_or_bytes, header=None)
        else:
            df = pd.read_excel(io.BytesIO(filepath_or_bytes), header=None)

        lines = []
        for _, row in df.iterrows():
            cells = [str(c).strip() for c in row if pd.notna(c) and str(c).strip()]
            if cells:
                lines.append(",".join(cells))
        text = "\n".join(lines)
    except Exception:
        return None

    if not text.strip():
        return None

    return load_candidate_text(text, filename)


# ============================================================
# Word文書 読み込み
# ============================================================

def load_candidate_docx(filepath_or_bytes, filename: str = "Word") -> Optional[Dict]:
    """Word(.docx)から候補者情報を抽出"""
    try:
        import docx
    except ImportError:
        return None

    try:
        if isinstance(filepath_or_bytes, (str, os.PathLike)):
            doc = docx.Document(filepath_or_bytes)
        else:
            doc = docx.Document(io.BytesIO(filepath_or_bytes))

        paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
        # テーブルからも抽出
        for table in doc.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    paragraphs.append(" ".join(cells))
        text = "\n".join(paragraphs)
    except Exception:
        return None

    if not text.strip():
        return None

    return load_candidate_text(text, filename)


# ============================================================
# 統合: ファイル形式を自動判定して読み込み
# ============================================================

SUPPORTED_EXTENSIONS = {
    ".csv": "CSV",
    ".txt": "テキスト",
    ".md": "テキスト",
    ".pdf": "PDF",
    ".png": "画像",
    ".jpg": "画像",
    ".jpeg": "画像",
    ".xlsx": "Excel",
    ".xls": "Excel",
    ".docx": "Word",
}


def load_candidate_file(filepath: str) -> Optional[Dict]:
    """ファイルパスから自動判定して候補者情報を読み込み"""
    ext = os.path.splitext(filepath)[1].lower()
    filename = os.path.basename(filepath)

    if ext == ".csv":
        return load_candidate_csv(filepath)
    elif ext in (".txt", ".md"):
        with open(filepath, "r", encoding="utf-8") as f:
            return load_candidate_text(f.read(), filename)
    elif ext == ".pdf":
        return load_candidate_pdf(filepath, filename)
    elif ext in (".png", ".jpg", ".jpeg"):
        return load_candidate_image(filepath, filename)
    elif ext in (".xlsx", ".xls"):
        return load_candidate_excel(filepath, filename)
    elif ext == ".docx":
        return load_candidate_docx(filepath, filename)
    return None


def load_candidate_upload(file_bytes: bytes, filename: str) -> Optional[Dict]:
    """Streamlitアップロードファイルから候補者情報を読み込み"""
    ext = os.path.splitext(filename)[1].lower()

    if ext == ".csv":
        text = file_bytes.decode("utf-8-sig")
        rows = list(csv.reader(io.StringIO(text)))
        # 既存CSV形式かチェック（候補者情報セクションがあるか）
        has_section = any(
            row and row[0].strip() in ("候補者情報", "項目", "候補者の強み", "強み")
            for row in rows
        )
        if has_section:
            candidate = {
                "name": filename,
                "info": {},
                "strengths": [],
                "keywords": [],
                "extra_keywords": [],
            }
            section = None
            full_lines = []
            for row in rows:
                if not row or all(cell.strip() == "" for cell in row):
                    continue
                first_cell = row[0].strip()
                full_lines.append(" ".join(c.strip() for c in row if c.strip()))
                if first_cell in ("候補者情報", "項目"):
                    section = "info"
                    continue
                elif first_cell in ("候補者の強み", "強み"):
                    section = "strengths"
                    continue
                elif first_cell in ("転職先候補求人リスト", "No."):
                    section = "jobs"
                    continue
                if section == "info" and len(row) >= 2:
                    key = row[0].strip()
                    val = row[1].strip()
                    if key and val and not _is_personal_info(key):
                        candidate["info"][key] = val
                elif section == "strengths" and len(row) >= 2:
                    sn = row[0].strip()
                    sd = row[1].strip() if len(row) > 1 else ""
                    if sn:
                        candidate["strengths"].append((sn, sd))
            full_text = "\n".join(full_lines)
            tags = extract_all_tags(full_text, candidate["info"])
            candidate["tags"] = tags
            candidate["conditions"] = _build_conditions(
                candidate["info"], candidate["strengths"], full_text, tags
            )
            return candidate
        else:
            return load_candidate_text(text, filename)

    elif ext in (".txt", ".md"):
        text = file_bytes.decode("utf-8-sig")
        return load_candidate_text(text, filename)

    elif ext == ".pdf":
        return load_candidate_pdf(file_bytes, filename)

    elif ext in (".png", ".jpg", ".jpeg"):
        return load_candidate_image(file_bytes, filename)

    elif ext in (".xlsx", ".xls"):
        return load_candidate_excel(file_bytes, filename)

    elif ext == ".docx":
        return load_candidate_docx(file_bytes, filename)

    return None


# ============================================================
# 複数ファイル統合
# ============================================================

def merge_candidate_uploads(files_data: List[Dict]) -> Dict:
    """
    複数ファイルから読み込んだ候補者データを統合する。
    履歴書 + 職務経歴書 + PF 等を1つのプロフィールにマージ。
    """
    if not files_data:
        return {}

    if len(files_data) == 1:
        return files_data[0]

    merged = {
        "name": "",
        "info": {},
        "strengths": [],
        "tags": {},
        "conditions": {},
        "source_files": [],
    }

    all_texts = []
    all_strengths = []
    all_info = {}

    for fd in files_data:
        if not fd:
            continue

        file_type = fd.get("_file_type", "その他")
        file_name = fd.get("name", "")
        merged["source_files"].append({
            "name": file_name,
            "type": file_type,
        })

        # 名前は最初に見つかったものを使用
        if not merged["name"] and fd.get("name"):
            merged["name"] = fd["name"]

        # info のマージ（後からのデータで上書き、ただし空でない場合のみ）
        for k, v in fd.get("info", {}).items():
            if v and (k not in all_info or not all_info[k]):
                all_info[k] = v

        # strengths のマージ（重複除外）
        existing_names = {s[0] for s in all_strengths}
        for s in fd.get("strengths", []):
            name = s[0] if isinstance(s, (list, tuple)) else s
            if name not in existing_names:
                all_strengths.append(s)
                existing_names.add(name)

        # テキスト全体を蓄積
        info_text = " ".join(str(v) for v in fd.get("info", {}).values())
        str_text = " ".join(
            f"{s[0]} {s[1]}" if isinstance(s, (list, tuple)) and len(s) > 1 else str(s)
            for s in fd.get("strengths", [])
        )
        all_texts.append(info_text + " " + str_text)

    merged["info"] = all_info
    merged["strengths"] = all_strengths

    # 統合テキストで全タグを再抽出
    combined_text = "\n".join(all_texts)
    merged["tags"] = extract_all_tags(combined_text, all_info)

    # 個別ファイルのタグもマージ
    for fd in files_data:
        if not fd or "tags" not in fd:
            continue
        fd_tags = fd["tags"]
        # リスト型タグのマージ
        for key in ["skills", "certifications", "industries", "employment_type",
                     "work_styles", "career_change_reasons", "interview_tags", "achievements"]:
            existing = set(merged["tags"].get(key, []))
            for item in fd_tags.get(key, []):
                if item not in existing:
                    merged["tags"].setdefault(key, []).append(item)
                    existing.add(item)

        # 語学のマージ
        existing_langs = {l["language"] for l in merged["tags"].get("languages", [])}
        for lang in fd_tags.get("languages", []):
            if lang["language"] not in existing_langs:
                merged["tags"].setdefault("languages", []).append(lang)
                existing_langs.add(lang["language"])

        # マネジメント経験（より詳細な方を優先）
        fd_mgmt = fd_tags.get("management", {})
        cur_mgmt = merged["tags"].get("management", {})
        if fd_mgmt.get("has_experience") and fd_mgmt.get("team_size", 0) > cur_mgmt.get("team_size", 0):
            merged["tags"]["management"] = fd_mgmt

        # 経験年数（より大きい方を採用）
        if fd_tags.get("experience_years", 0) > merged["tags"].get("experience_years", 0):
            merged["tags"]["experience_years"] = fd_tags["experience_years"]

        # 職歴社数（より多い方）
        if fd_tags.get("job_history_count", 0) > merged["tags"].get("job_history_count", 0):
            merged["tags"]["job_history_count"] = fd_tags["job_history_count"]

        # 入社可能時期（空でなければ上書き）
        if fd_tags.get("availability") and not merged["tags"].get("availability"):
            merged["tags"]["availability"] = fd_tags["availability"]

    # 経験レベルを統合情報で再判定
    merged["tags"]["experience_level"] = _extract_experience_level(
        combined_text, merged["tags"].get("experience_years", 0)
    )

    # conditionsを統合タグベースで再構築
    merged["conditions"] = _build_conditions(
        merged["info"], merged["strengths"], combined_text, merged["tags"]
    )

    return merged


# ============================================================
# 既存CSVファイルの一括読み込み（後方互換）
# ============================================================

def list_candidate_csvs() -> List[str]:
    """候補者CSVファイルの一覧を取得"""
    pattern = os.path.join(CSV_DIR, "[0-9][0-9]_候補者*.csv")
    return sorted(glob.glob(pattern))


def get_candidate_display_name(filepath: str) -> str:
    """ファイル名から表示用の候補者名を取得"""
    basename = os.path.basename(filepath)
    match = re.match(r'\d+_(候補者\d+)_(\d+歳)_(.+)\.csv', basename)
    if match:
        return f"{match.group(1)} ({match.group(2)}・{match.group(3)})"
    # 拡張子を除いたファイル名
    name = os.path.splitext(basename)[0]
    # 番号プレフィックスを除去
    name = re.sub(r'^\d+[_\-]', '', name)
    return name


def load_all_candidates() -> List[Dict]:
    """全候補者CSVを読み込み"""
    csvs = list_candidate_csvs()
    candidates = []
    for path in csvs:
        cand = load_candidate_csv(path)
        if cand:
            cand["display_name"] = get_candidate_display_name(path)
            candidates.append(cand)
    return candidates
